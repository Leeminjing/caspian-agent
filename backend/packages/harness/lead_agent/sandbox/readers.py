"""
本文件提供文档格式解析器，从 LocalSandbox 中解耦为独立模块。

对外提供:
    _read_pdf(real_path)   : .pdf  文件文字提取（含扫描版 OCR + 版面解析）
    _read_docx(real_path)  : .docx 文件文字提取（段落 + 表格）
    _read_doc(real_path)   : .doc  文件文字提取（olefile → LibreOffice fallback）

输入:
    所有函数仅接受已解析的真实磁盘路径 (real_path: str)，返回提取的文字内容 (str)。
    路径解析与安全校验由调用方（LocalSandbox.read_file）在调用前完成。

工作流:
    _read_pdf:
    (1) pypdf 逐页提取文字
    (2) 某页为空 → pypdfium2 渲染为图像 → pytesseract OCR → _reconstruct_layout 版面重建
    (3) Tesseract 未安装时跳过 OCR，返回 pypdf 结果

    _read_docx:
    (1) python-docx 逐段提取文字
    (2) 表格逐行逐格提取，单元格以制表符分隔，行间以换行符分隔

    _read_doc:
    (1) olefile 从 WordDocument 流提取 Unicode 文本
    (2) 失败则尝试 soffice --headless --convert-to txt
    (3) 两者均不可用则抛 RuntimeError

    _reconstruct_layout:
    (1) 按 block_num 分组 → 多栏检测（水平间距 > 页面宽度 20%）
    (2) 排序：上→下优先，同栏左→右
    (3) block 内按 par_num → line_num → word_num 展开，段落间空行

解耦约束:
    本模块不 import local.py / base.py 等沙箱模块。
    函数为纯输入输出，无沙箱状态依赖。

示例:
    text = _read_pdf("/path/to/report.pdf")
    text = _read_docx("/path/to/memo.docx")
    text = _read_doc("/path/to/legacy.doc")
"""

import os
import shutil
import subprocess
import tempfile

from pypdf import PdfReader


def _read_pdf(real_path: str) -> str:
    pages = []
    reader = PdfReader(real_path)
    for i, page in enumerate(reader.pages):
        text = page.extract_text()
        if text and text.strip():
            pages.append(text.strip())
        else:
            ocr_text = _ocr_pdf_page(real_path, i)
            pages.append(ocr_text.strip() if ocr_text else "")
    return "\n".join(pages)


def _ocr_pdf_page(real_path: str, page_index: int) -> str:
    """使用 pypdfium2 渲染指定页为图像，再通过 pytesseract OCR 识别。"""
    if shutil.which("tesseract") is None:
        return ""

    try:
        import pypdfium2 as pdfium
        import pytesseract
    except ImportError:
        return ""

    pdf_doc = pdfium.PdfDocument(real_path)
    page = pdf_doc[page_index]
    bitmap = page.render(scale=2)
    image = bitmap.to_pil()
    pdf_doc.close()

    try:
        data = pytesseract.image_to_data(
            image,
            output_type=pytesseract.Output.DICT,
            lang="chi_sim+eng",
        )
        return _reconstruct_layout(data)
    except Exception:
        return ""


def _read_docx(real_path: str) -> str:
    from docx import Document

    doc = Document(real_path)
    parts = []

    for para in doc.paragraphs:
        parts.append(para.text)

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text for cell in row.cells]
            parts.append("\t".join(cells))
        parts.append("")

    return "\n".join(parts)


def _read_doc(real_path: str) -> str:
    try:
        return _extract_doc_olefile(real_path)
    except Exception:
        pass

    if shutil.which("soffice"):
        try:
            return _extract_doc_libreoffice(real_path)
        except Exception:
            pass

    raise RuntimeError(
        "无法读取 .doc 文件。请安装 LibreOffice 或将文件转换为 .docx 格式。"
    )


def _extract_doc_olefile(real_path: str) -> str:
    import olefile

    ole = olefile.OleFileIO(real_path)
    if not ole.exists("WordDocument"):
        ole.close()
        raise ValueError("Not a valid .doc OLE2 file (no WordDocument stream)")

    text = ""
    for stream_name in ole.listdir():
        name = "/".join(stream_name)
        if name == "WordDocument":
            continue
        try:
            data = ole.openstream(stream_name).read()
            text += _extract_unicode_from_ole(data)
        except Exception:
            continue

    ole.close()
    return text.strip() or _extract_raw_text_from_ole(real_path)


def _extract_unicode_from_ole(data: bytes) -> str:
    result = []
    i = 0
    while i < len(data) - 1:
        char = data[i:i + 2]
        i += 2
        if char == b"\x00\x00":
            continue
        try:
            decoded = char.decode("utf-16-le")
            if decoded.isprintable() or decoded in ("\n", "\r", "\t", " "):
                result.append(decoded)
        except UnicodeDecodeError:
            continue
    return "".join(result)


def _extract_raw_text_from_ole(real_path: str) -> str:
    import olefile

    ole = olefile.OleFileIO(real_path)
    text = ""
    word_stream = ole.openstream("WordDocument").read()
    text += _extract_unicode_from_ole(word_stream)
    ole.close()
    return text


def _extract_doc_libreoffice(real_path: str) -> str:
    with tempfile.TemporaryDirectory() as tmpdir:
        subprocess.run(
            ["soffice", "--headless", "--convert-to", "txt", "--outdir", tmpdir, real_path],
            check=True,
            capture_output=True,
            timeout=60,
        )
        txt_name = os.path.splitext(os.path.basename(real_path))[0] + ".txt"
        txt_path = os.path.join(tmpdir, txt_name)
        with open(txt_path, "r", encoding="utf-8") as f:
            return f.read()


def _reconstruct_layout(data: dict) -> str:
    words = []
    n = len(data["text"])
    page_right = max(
        (data["left"][i] + data["width"][i])
        for i in range(n)
        if data["text"][i].strip()
    ) or 1

    for i in range(n):
        word = data["text"][i].strip()
        if not word:
            continue
        words.append({
            "block": data["block_num"][i],
            "par": data["par_num"][i],
            "line": data["line_num"][i],
            "word": data["word_num"][i],
            "left": data["left"][i],
            "top": data["top"][i],
            "width": data["width"][i],
            "height": data["height"][i],
            "text": word,
        })

    if not words:
        return ""

    blocks = {}
    for w in words:
        blocks.setdefault(w["block"], []).append(w)

    sorted_blocks = _sort_blocks_by_reading_order(blocks, page_right)

    result = []
    for block_words, _column in sorted_blocks:
        pars = {}
        for w in block_words:
            pars.setdefault(w["par"], []).append(w)

        for par_num in sorted(pars):
            lines = {}
            for w in pars[par_num]:
                lines.setdefault(w["line"], []).append(w)

            for line_num in sorted(lines):
                line_words = sorted(lines[line_num], key=lambda w: w["word"])
                result.append(" ".join(w["text"] for w in line_words))

            result.append("")

    return "\n".join(result).strip()


def _sort_blocks_by_reading_order(blocks: dict, page_right: int) -> list:
    block_info = []
    for bid, words in blocks.items():
        tops = [w["top"] for w in words]
        lefts = [w["left"] for w in words]
        block_info.append({
            "id": bid,
            "words": words,
            "top": min(tops),
            "left": min(lefts),
            "bottom": max([t + words[i]["height"] for i, t in enumerate(tops)]),
        })

    COLUMN_GAP_RATIO = 0.2
    column_groups = _detect_columns(block_info, page_right, COLUMN_GAP_RATIO)

    result = []
    for column_idx, col_blocks in enumerate(column_groups):
        col_blocks.sort(key=lambda b: b["top"])
        for b in col_blocks:
            result.append((b["words"], column_idx))

    return result


def _detect_columns(block_info: list, page_right: int, gap_ratio: float) -> list:
    if not block_info:
        return [[]]

    gap_threshold = page_right * gap_ratio
    sorted_by_left = sorted(block_info, key=lambda b: b["left"])

    columns = []
    current_col = [sorted_by_left[0]]
    current_col_right = sorted_by_left[0]["left"]

    for b in sorted_by_left[1:]:
        for prev in current_col:
            prev_right = prev["left"] + max(
                w["left"] + w["width"] - prev["left"]
                for w in prev["words"]
            )
            current_col_right = max(current_col_right, prev_right)

        if b["left"] - current_col_right > gap_threshold:
            columns.append(current_col)
            current_col = [b]
            current_col_right = b["left"]
        else:
            current_col.append(b)

    columns.append(current_col)
    return columns
